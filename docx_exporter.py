"""DOCX exporter for ByteCase Playbooks."""
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.shared import Pt

from playbook_data import APP_NAME, APP_VERSION, APP_ATTRIBUTION, APP_DOMAIN, PLAYBOOK_BOUNDARY, get_playbook
from session_core import session_summary


def add_bullets(doc: Document, values):
    for value in values or []:
        doc.add_paragraph(str(value), style="List Bullet")


def export_session_docx(session: Dict[str, Any], output_path) -> Path:
    output_path = Path(output_path)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10)

    doc.add_heading(APP_NAME, level=0)
    doc.add_paragraph("Guided Examiner Reference Session")
    doc.add_paragraph(APP_ATTRIBUTION)
    doc.add_paragraph(APP_DOMAIN)

    meta = doc.add_table(rows=0, cols=2)
    for label, value in [
        ("Version", APP_VERSION),
        ("Mode", session.get("mode", "")),
        ("Playbook", session.get("playbook_title", "")),
        ("Current Step", session.get("current_step_index", 1)),
        ("Created", session.get("created_at", "")),
        ("Updated", session.get("updated_at", "")),
    ]:
        row = meta.add_row().cells
        row[0].text = label
        row[1].text = str(value or "")

    doc.add_heading("Boundary Notice", level=1)
    doc.add_paragraph(PLAYBOOK_BOUNDARY)

    playbook = get_playbook(session.get("playbook_id", "")) or {}
    if playbook:
        doc.add_heading("Playbook Summary", level=1)
        doc.add_paragraph(playbook.get("summary", ""))
        doc.add_heading("Use When", level=2)
        add_bullets(doc, playbook.get("use_when", []))
        doc.add_heading("Avoid / Pause When", level=2)
        add_bullets(doc, playbook.get("avoid_when", []))

        state = {item.get("index"): item for item in session.get("step_state", [])}
        current_step = session.get("current_step_index", 1)
        doc.add_heading("Steps", level=1)
        for idx, step in enumerate(playbook.get("steps", []), start=1):
            item = state.get(idx, {})
            mark = "[x]" if item.get("checked") else "[ ]"
            pointer = " - current step" if idx == current_step else ""
            doc.add_heading(f"{mark} Step {idx}: {step.get('title', '')}{pointer}", level=2)
            doc.add_paragraph(f"Field Focus: {step.get('field_focus', '')}")
            doc.add_paragraph(f"Learning Detail: {step.get('learning_detail', '')}")
            doc.add_paragraph(f"Why: {step.get('why', '')}")
            for label, key in [("Possible Tools", "tools"), ("Common Artifacts", "artifacts"), ("Cautions", "cautions"), ("Does Not Prove", "does_not_prove"), ("Document", "document")]:
                values = step.get(key, [])
                if values:
                    doc.add_paragraph(label + ":")
                    add_bullets(doc, values)
            commands = step.get("command_examples", [])
            if commands:
                doc.add_paragraph("Command Examples:")
                for command in commands:
                    doc.add_paragraph(f"{command.get('label', 'Command')}: {command.get('example', '')}", style="List Bullet")
                    doc.add_paragraph(f"Purpose: {command.get('purpose', '')}")
                    doc.add_paragraph(f"Does not prove: {command.get('does_not_prove', '')}")
            notes = (item.get("notes") or "").strip()
            if notes:
                doc.add_paragraph("Step Reference Notes:")
                doc.add_paragraph(notes)

    summary = session_summary(session)
    doc.add_heading("Session Summary", level=1)
    doc.add_paragraph(f"Current Step: {summary['current_step_index']} of {summary['total_steps']}")
    doc.add_paragraph(f"Reviewed Steps: {summary['checked_steps']} of {summary['total_steps']}")
    doc.add_paragraph(f"Steps With Notes: {summary['steps_with_notes']}")

    doc.save(output_path)
    return output_path
